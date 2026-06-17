# -*- coding: utf-8 -*-
"""
HIA 评测语料定向采集器 —— 基于 gov_bumen_crawler.py 改造(自包含,可移植)
来源: https://sousuo.www.gov.cn/zcwjk/policyDocumentLibrary  (t=zhengcelibrary_bm)

相对原脚本(全部门按配额广撒网 1000 份)的 5 处定向改造,均为对齐 HIA 采样清单:
  ① OUT_DIR 改成易改的输出目录(另一台电脑只改这一行)。
  ② WHITELIST: 只抓采样清单要的部门(卫健/生态环境/住建/市场监管/交通/民政/农业/教育/人社),
     覆盖健康维度 1/3/5/6/7/9/10 + 跨部门"含健康词的负样本"。
  ③ ATT_EXT 只收 .pdf/.docx(引擎只支持这两种;原脚本的 doc/xls/wps/ofd 抽不了文本)。
  ④ 每部门固定配额 QUOTA_PER_DEPT,不再按全国文件占比分配(否则财政/税务等大部门吃光名额)。
  ⑤ 保留:连通校验 / 断点续传(progress.json) / 死循环保护 / index.xlsx 索引。

依赖:  pip install requests beautifulsoup4 lxml python-docx openpyxl tqdm
用法:  python hia_policy_crawler.py            # 按下方配置抓取
       python hia_policy_crawler.py 8          # 覆盖每部门配额为 8
"""

import os
import re
import time
import json
import random
import sys
import requests
from urllib.parse import urljoin

# Windows 控制台/重定向默认 GBK,政策标题含   等字符会让 print 崩溃 → 强制 UTF-8
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import Workbook
from tqdm import tqdm

# ============================ 配置(改这里)============================
# ① 输出目录 —— 另一台电脑改成你的路径即可
OUT_DIR = r"E:\projects\test2"

# ② 部门白名单(对齐 HIA 采样清单;名称须与 gov.cn facet 完全一致)
WHITELIST = [
    "国家卫生健康委员会",     # 维度1传染病/2慢病/4突发公卫/5人口/9优质医疗/10服务可及 + 程序类负样本
    "生态环境部",            # 维度6:噪声/水/土壤/大气(触发非空气类证据卡)
    "住房和城乡建设部",       # 维度7:绿地/人居/城市更新
    "国家市场监督管理总局",   # 维度3/6:食品安全
    "交通运输部",            # 维度3/7:道路安全/慢行系统
    "民政部",               # 维度5:养老/殡葬
    "农业农村部",            # 维度6:人居环境/农产品
    "教育部",               # 学校卫生 + 大量程序类负样本
    "人力资源和社会保障部",    # 职业健康/工伤 + 程序类负样本
]

# ④ 每部门配额(有附件/无附件各占一半);白名单 9 部门 × 默认 6 ≈ 54 份
QUOTA_PER_DEPT = int(sys.argv[1]) if len(sys.argv) > 1 else 6

# 大集模式:env CRAWL_ALL=1 → 实时拉 gov.cn 全部门 facet,自动排除下面已抓的 9 个,
# 抓其余所有部门(广覆盖)。默认关闭,保持 9 部门 WHITELIST 原行为。
CRAWL_ALL = os.environ.get("CRAWL_ALL", "") not in ("", "0", "false", "False")
DONE_DEPTS = set(WHITELIST)   # 已抓过、CRAWL_ALL 时跳过的部门(= 上面 9 部门)

# ③ 只收引擎支持的格式
ATT_EXT = (".pdf", ".docx")

SLEEP        = (0.6, 1.4)     # 每请求节流,礼貌爬
MAX_PAGE     = 60            # 单部门最多翻页,死循环保护(1000份大集:30→60够翻完大部门)
NO_NEW_BREAK = 4            # 连续 N 页无新增就停
BALANCE_ATT_TXT = False      # 1000份大集:关掉"有附件/无附件各半"内部配额,防大部门凑不满
TEXT_FALLBACK_MINLEN = 120   # 无附件时,正文≥此长度才转 docx
# ===================================================================

API  = "https://sousuo.www.gov.cn/search-gov/data"
T    = "zhengcelibrary_bm"
TYPE = "gwyzcwjk"
INDEX_XLSX = os.path.join(OUT_DIR, "_index.xlsx")
PROGRESS   = os.path.join(OUT_DIR, "_progress.json")

HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                   "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"),
    "Referer": "https://sousuo.www.gov.cn/zcwjk/policyDocumentLibrary",
}
session = requests.Session()
session.headers.update(HEADERS)


def nap():
    time.sleep(random.uniform(*SLEEP))


def safe_name(s, maxlen=70):
    s = re.sub(r'[\\/:*?"<>|\n\r\t]', "_", s or "untitled")
    s = re.sub(r"\s+", " ", s).strip()
    return s[:maxlen]


def fetch_all_departments():
    """从 gov.cn bmfl facet 实时拉全部门名(按文件数降序)。"""
    params = dict(t=T, q="", sort="pubtime", sortType="1", searchfield="title",
                  p="1", n="1", type=TYPE)
    r = session.get(API, params=params, timeout=30)
    r.raise_for_status()
    bmfl = r.json().get("searchVO", {}).get("extendresult", {}) \
            .get("facetMap", {}).get("bmfl", {}) or {}
    def num(v):
        m = re.search(r"\d+", str(v))
        return int(m.group()) if m else 0
    depts = sorted(((k, num(v)) for k, v in bmfl.items() if k != "count"),
                   key=lambda x: -x[1])
    return [k for k, _ in depts]


def fetch_list_page(dept, page, n=50):
    """抓某部门第 page 页列表(用 bmfl 筛,与 facet 同源)。"""
    params = dict(t=T, q="", timetype="", mintime="", maxtime="",
                  sort="pubtime", sortType="1", searchfield="title",
                  bmfl=dept or "", puborg="", pcodeYear="", pcodeNum="",
                  filetype="", p=str(page), n=str(n),
                  inpro="", dup="", orpro="", type=TYPE)
    r = session.get(API, params=params, timeout=30)
    r.raise_for_status()
    sv = r.json().get("searchVO", {}) or {}
    return sv.get("listVO", []) or [], sv.get("totalCount", 0)


def parse_detail(url):
    """访问详情页,返回 (meta, 正文文本, 合规附件URL列表)。"""
    r = session.get(url, timeout=30)
    r.raise_for_status()
    r.encoding = r.apparent_encoding or "utf-8"
    # gov.cn 详情页 HTML 不规范:lxml 会提前截断 DOM、丢掉正文容器(UCAP-CONTENT/pages_content);
    # html.parser 更宽容,能正确解析。这是脚本跑通这个数据源的关键。
    soup = BeautifulSoup(r.text, "html.parser")

    meta = {}
    full_text = soup.get_text(" ", strip=True)
    patterns = {
        "发文机关": r"发文机关[：:\s]*([^\n]+?)(?:发文字号|来\s*源|主题分类|公文种类|成文日期|$)",
        "发文字号": r"发文字号[：:\s]*([^\n]+?)(?:来\s*源|主题分类|公文种类|成文日期|发文机关|$)",
        "来源":     r"来\s*源[：:\s]*([^\n]+?)(?:主题分类|公文种类|成文日期|发文字号|$)",
        "主题分类": r"主题分类[：:\s]*([^\n]+?)(?:公文种类|成文日期|来\s*源|$)",
        "公文种类": r"公文种类[：:\s]*([^\n]+?)(?:成文日期|主题分类|$)",
        "成文日期": r"成文日期[：:\s]*([^\n]+?)(?:\s{2,}|$)",
    }
    for k, pat in patterns.items():
        m = re.search(pat, full_text)
        meta[k] = m.group(1).strip() if m else ""

    cont = (soup.select_one(".pages_content") or soup.select_one(".TRS_Editor")
            or soup.select_one("#UCAP-CONTENT") or soup.select_one(".article"))
    body_text = cont.get_text("\n", strip=True) if cont else ""

    atts = []
    for a in soup.select("a[href]"):
        href = a.get("href", "")
        if href and href.lower().split("?")[0].endswith(ATT_EXT):
            atts.append(urljoin(url, href))
    return meta, body_text, list(dict.fromkeys(atts))


def download_attachment(att_url, base_name, idx):
    ext = os.path.splitext(att_url.split("?")[0])[1] or ".bin"
    fname = f"{base_name}{'' if idx == 0 else f'_{idx}'}{ext}"
    fpath = os.path.join(OUT_DIR, fname)
    with session.get(att_url, timeout=120, stream=True) as r:
        r.raise_for_status()
        with open(fpath, "wb") as f:
            for chunk in r.iter_content(8192):
                f.write(chunk)
    return os.path.basename(fpath)


def save_text_as_docx(title, body_text, meta, base_name):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in [f"发文机关：{meta.get('发文机关','')}",
                 f"发文字号：{meta.get('发文字号','')}",
                 f"公文种类：{meta.get('公文种类','')}",
                 f"主题分类：{meta.get('主题分类','')}",
                 f"成文日期：{meta.get('成文日期','')}",
                 f"来源：{meta.get('来源','')}"]:
        doc.add_paragraph(line)
    doc.add_paragraph("—" * 30)
    for para in body_text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    fpath = os.path.join(OUT_DIR, base_name + ".docx")
    doc.save(fpath)
    return os.path.basename(fpath)


def load_progress():
    if os.path.exists(PROGRESS):
        with open(PROGRESS, encoding="utf-8") as f:
            return json.load(f)
    return {"seen_urls": [], "rows": []}


def save_progress(seen_urls, rows):
    with open(PROGRESS, "w", encoding="utf-8") as f:
        json.dump({"seen_urls": list(seen_urls), "rows": rows},
                  f, ensure_ascii=False, indent=2)


def write_index(rows):
    wb = Workbook(); ws = wb.active; ws.title = "索引"
    ws.append(["序号", "发布部门", "公文种类", "主题分类", "标题",
               "发文字号", "成文日期", "处理方式", "本地文件", "来源URL"])
    for row in rows:
        ws.append(row)
    wb.save(INDEX_XLSX)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    if CRAWL_ALL:
        whitelist = [d for d in fetch_all_departments() if d not in DONE_DEPTS]
        print(f"[CRAWL_ALL] 全部门模式:抓 {len(whitelist)} 个新部门"
              f"(已排除 {len(DONE_DEPTS)} 个已抓),每部门约 {QUOTA_PER_DEPT} 份 → {OUT_DIR}\n")
    else:
        whitelist = WHITELIST
        print(f"定向采集 {len(whitelist)} 个部门,每部门约 {QUOTA_PER_DEPT} 份 → {OUT_DIR}\n")

    state = load_progress()
    seen_urls = set(state["seen_urls"])
    rows = state["rows"]
    count = len(rows)

    for dept in whitelist:
        got, page, no_new = 0, 1, 0
        half = max(1, QUOTA_PER_DEPT // 2)
        if BALANCE_ATT_TXT:
            cap_att, cap_txt = half, QUOTA_PER_DEPT - half
        else:
            cap_att = cap_txt = QUOTA_PER_DEPT
        got_att, got_txt = 0, 0
        while got < QUOTA_PER_DEPT and page <= MAX_PAGE:
            try:
                items, _ = fetch_list_page(dept, page); nap()
            except Exception as e:                               # noqa: BLE001
                print(f"  [列表失败] {dept} p{page}: {e}"); break
            if not items:
                break
            page_new = 0
            for it in items:
                if got >= QUOTA_PER_DEPT:
                    break
                url = it.get("url", "")
                if not url or url in seen_urls:
                    continue
                seen_urls.add(url)
                title = re.sub(r"<[^>]+>", "", it.get("title", "")).strip()
                base = f"{safe_name(dept,6)}_{count+1:03d}_{safe_name(title)}"
                try:
                    meta, body, atts = parse_detail(url); nap()
                except Exception as e:                           # noqa: BLE001
                    print(f"  [详情失败] {url}: {e}"); continue
                has_att = bool(atts)
                # 二级配额:有/无附件各半,某类满了就把名额让给另一类
                if has_att and got_att >= cap_att and got_txt < cap_txt:
                    continue
                if (not has_att) and got_txt >= cap_txt and got_att < cap_att:
                    continue
                files, mode = [], ""
                try:
                    if has_att:
                        for i, a in enumerate(atts):
                            files.append(download_attachment(a, base, i)); nap()
                        mode = "附件"
                    elif len(body) >= TEXT_FALLBACK_MINLEN:
                        files.append(save_text_as_docx(title, body, meta, base))
                        mode = "正文转docx"
                    else:
                        continue
                except Exception as e:                           # noqa: BLE001
                    print(f"  [保存失败] {url}: {e}"); continue
                if not files:
                    continue
                count += 1; got += 1; page_new += 1
                if has_att: got_att += 1
                else:       got_txt += 1
                rows.append([count, meta.get("发文机关") or dept,
                             meta.get("公文种类", ""), meta.get("主题分类", ""),
                             title, meta.get("发文字号", ""), meta.get("成文日期", ""),
                             mode, "; ".join(files), url])
                print(f"  [{count}] {dept[:6]} | {mode} | {title[:38]}")
                if count % 10 == 0:
                    save_progress(seen_urls, rows); write_index(rows)
            no_new = no_new + 1 if page_new == 0 else 0
            if no_new >= NO_NEW_BREAK:
                break
            page += 1
        print(f"  → {dept} 抓到 {got} 份\n")

    save_progress(seen_urls, rows); write_index(rows)
    print(f"完成!共 {count} 份 → {OUT_DIR}  (索引 {INDEX_XLSX})")


if __name__ == "__main__":
    main()
