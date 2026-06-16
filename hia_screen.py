# -*- coding: utf-8 -*-
"""AI 辅助 HIA 初筛 —— 引擎(因果路径版)。

对标《健康影响评估初筛表》。核心不是"文档→直接判 10 题",而是显式建出
**政策行动 → 健康决定因素(多级、间接)→ 健康结果(10 题)** 的因果路径网,
再由代码按确定性阈值聚合到 10 题。让间接/多级路径被系统化展开,贴合 HIA 的
logic framework 与社会健康决定因素(SDH/Dahlgren–Whitehead)方法。

流水线(被 views/hia_screen.py 调用):
  ① extract_actions   文档 → 政策行动/要素
  ② expand_pathways   行动 → 多视角(环境/社会心理/公平/卫生系统)因果路径(深度2–3)
  ③ critique_augment  完整性批判:补漏的决定因素/脆弱群体/间接效应 + 小结 + 程度建议
  · compute_items     代码确定性聚合到 10 题(判断阈值不交给 LLM)
  · build_dot         生成可渲染的因果路径图(DOT,st.graphviz_chart 用)
AI 仅辅助研判与展开路径,采纳/剪枝/判定/签字以专家为准。
"""
import json
import re
from io import BytesIO

import requests

import hia_evidence

API_URL = "https://api.deepseek.com/chat/completions"
MODEL = "deepseek-chat"
MAX_DOC_CHARS = 40000
ANSWERS = ("是", "不知道", "否")
STRENGTHS = ("强", "中", "推测")
STATUSES = ("文档支持", "路径库/文献", "假设待证")

QUESTIONS = [
    "可能导致人群传染病和感染性疾病的发生发展。",
    "可能加剧人群重点慢性病的发生发展。",
    "可能增加人群中毒和伤害事件的风险。",
    "可能增加其他突发公共卫生事件的风险。",
    "可能对人口高质量发展带来不利影响。",
    "可能对空气、饮用水、食品和环境卫生等健康环境带来不利影响。",
    "可能对人群健康生活方式、社会心理健康等带来不利影响。",
    "可能对卫生健康投入保障和医疗保险水平带来不利影响。",
    "可能对优质医疗资源合理配置和利用带来不利影响。",
    "可能对医疗卫生服务质量安全和利用、公平性和可及性带来不利影响。",
]
# 图/UI 用的短标签
SHORT_Q = ["传染病", "重点慢病", "中毒伤害", "突发公卫", "人口发展",
           "健康环境", "生活方式/心理", "卫生投入/医保", "优质医疗资源", "服务质量/可及"]

# —— 健康决定因素脚手架(grounding;让路径展开系统化而非随性联想)——
DETERMINANTS = """健康决定因素清单(展开路径时逐层套用,可多级串联):
- 个体生活方式:体力活动、饮食、吸烟饮酒、出行方式选择
- 社会与社区:社会支持/凝聚力、社会心理压力、孤独、社会资本
- 生活与工作条件:住房、交通与道路安全、职业暴露与工作环境、就业与收入、教育、
  食品供应与安全、给排水与环境卫生、医疗卫生服务可及性
- 环境要素:空气质量(PM2.5/NO2)、水质、噪声、热环境/绿地、土壤、危化品
- 总体社会经济/文化/环境:人口结构、经济结构、城市空间形态、健康公平"""

LENSES = ["环境健康", "社会心理与行为", "健康公平与脆弱群体", "卫生系统与服务"]


# ============ 文档抽取 ============
def extract_text(name, data):
    info = {"kind": "", "pages": 0, "truncated": False, "error": ""}
    low = (name or "").lower()
    try:
        if low.endswith(".pdf"):
            info["kind"] = "PDF"
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(data))
            info["pages"] = len(reader.pages)
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
        elif low.endswith(".docx"):
            info["kind"] = "Word"
            from docx import Document
            doc = Document(BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append("\t".join(c.text for c in row.cells))
            text = "\n".join(parts)
        elif low.endswith(".doc"):
            info["error"] = "旧版 .doc 不支持,请另存为 .docx 或 PDF 后上传。"
            return "", info
        else:
            info["error"] = "仅支持 PDF 或 Word(.docx)。"
            return "", info
    except Exception as e:
        info["error"] = f"文档解析失败:{e}"
        return "", info
    text = (text or "").strip()
    if not text:
        info["error"] = "未能从文档提取到文字(可能是扫描件/图片型 PDF,需 OCR;本工具暂不支持)。"
    if len(text) > MAX_DOC_CHARS:
        text = text[:MAX_DOC_CHARS]
        info["truncated"] = True
    return text, info


def guess_title(doc_text, fallback=""):
    """从政策正文猜测评估对象标题:优先「(关于)印发《X》」中的 X(即被印发的政策本身),
    其次取正文前段最长的《…》,都没有则回退到 fallback(通常是文件名)。仅做确定性提取。"""
    t = doc_text or ""
    # 书名号在 PDF 抽取后可能变体为 «» / 〈〉,一并兼容。
    m = re.search(r"印发\s*[《«〈](.+?)[》»〉]", t)
    if m:
        return m.group(1).strip()
    titles = re.findall(r"[《«〈](.+?)[》»〉]", t[:1500])
    if titles:
        return max(titles, key=len).strip()
    return (fallback or "").strip()


# ============ LLM 调用 ============
def _extract_json(content):
    s = (content or "").strip()
    if s.startswith("```"):
        s = s.strip("`").strip()
        if s[:4].lower() == "json":
            s = s[4:].strip()
    i, j = s.find("{"), s.rfind("}")
    if i >= 0 and j > i:
        s = s[i:j + 1]
    try:
        d = json.loads(s)
        return d if isinstance(d, dict) else {}
    except Exception:
        return {}


def _chat_json(system, user, key, timeout=120, max_tokens=4000, temps=(0.3, 0.8)):
    """调 DeepSeek(json_object)+ 稳健提取 + 升温重试,返回 dict(失败为 {})。"""
    for temp in temps:
        try:
            r = requests.post(
                API_URL,
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                json={"model": MODEL,
                      "messages": [{"role": "system", "content": system},
                                   {"role": "user", "content": user}],
                      "temperature": temp, "response_format": {"type": "json_object"},
                      "max_tokens": max_tokens},
                timeout=timeout)
            r.raise_for_status()
            data = _extract_json(r.json()["choices"][0]["message"]["content"] or "")
            if data:
                return data
        except Exception:
            continue
    return {}


def _q_block():
    return "\n".join(f"{i+1}. {q}" for i, q in enumerate(QUESTIONS))


# ============ ① 行动抽取 ============
_SYS_ACTIONS = """你是健康影响评估(HIA)助手。阅读"评估对象"文档,只抽取该政策/规划/工程
**实际引入、文档中确有依据的具体行动/要素/改变**(建设、用地、交通、产业、环境、设施、
能源、人口就业等,无论是否直接关乎健康)。每条尽量具体、可核验,并给文档原文依据。

先判**是否健康相关**(health_relevant):
- **false**:文档主体是**表彰/奖励/任免/命名/纪念/会议/防空警报试鸣/财税分担/行政处罚程序/
  废止/统计**等、**不引入任何实体、环境或行为改变**的内容(此类与人群健康关联很弱)。
- **true**:涉及**建设/用地/交通/产业/环境/能源/住房/生态/卫生健康/人口就业**等可能影响人群
  健康的实质内容。

铁律(防止臆造):
- **只抽文档真实写到的**;绝不臆造文档没有的事实;**绝不把本提示里的任何措辞当作行动照抄**。
- 健康相关(true)时,抽出文档里的主要行动(具体、可核验、带原文依据);宁少勿凑。
- 非健康(false)时,actions 返回**空列表**。

只输出 JSON:{"health_relevant": true 或 false, "policy_type":"政策类型一词",
 "actions":[{"id":"A1","action":"简明行动描述","evidence":"文档原文依据"}]}
id 用 A1、A2…。"""


def extract_actions(doc_text, key, project_name=""):
    """返回 (actions, health_relevant)。health_relevant=False → 该政策与健康关联很弱,上层据此门控。"""
    user = f"【评估对象文档(可能已截断)】\n{doc_text}\n\n请先判健康相关性再抽取行动,只输出 JSON。"
    data = _chat_json(_SYS_ACTIONS, user, key, max_tokens=2500)
    out = []
    for i, a in enumerate(data.get("actions") or [], 1):
        if not isinstance(a, dict):
            continue
        act = str(a.get("action", "") or "").strip()
        if not act:
            continue
        out.append({"id": f"A{i}", "action": act,
                    "evidence": str(a.get("evidence", "") or "").strip()})
    # 仅当模型显式判 false 才视为非健康(缺省/异常一律按 True,避免误伤)
    health_relevant = data.get("health_relevant", True) is not False
    # 安全网:模型偶发把"绿色低碳/节能减排/生态环境"等误判为非健康 → 标题(文件名)含强关键词则纠回 True
    if not health_relevant:
        title = (project_name or "") + " " + (doc_text or "")[:400]
        if re.search(r"绿色|低碳|减排|节能|生态|环境|污染|排放|气候|卫生|健康|医疗|养老|住房|"
                     r"交通|能源|绿地|噪声|食品|垃圾|废物|供水|饮用水|城镇化|城市更新|"
                     r"规划|建设|改造|整治", title):
            health_relevant = True
    return out, health_relevant


# ============ ② 路径展开 ============
_SYS_EXPAND = f"""你是 HIA 因果路径分析专家。给定政策"行动"清单,系统化地推演每个行动如何
**经由健康决定因素、多级且常常间接地**影响人群健康,最终落到初筛表的某一个重点问题上。

务必先判健康关联(**看政策目标,不要孤立看单条行动的字面**):
0a. **服务于减排/降碳/节能/清洁能源/生态/污染治理/绿色低碳目标的行动,即使单条是工程技术
    指标也要展开**:像"供电煤耗270克/千瓦时""机组负荷变化速率""度电碳排放降10%""储能配置"
    "调峰能力"这类单看是纯参数,但它们整体服务于减污降碳——应把同目标的技术行动**归并**,
    展开该政策"减排/降污 → 空气/水/温室气体改善 → 人群健康"的路径。**不要因为单条像技术
    参数就跳过,否则会漏掉煤电升级、储能、节能降碳这类核心环境健康政策。**
0b. **纯经济/管理类政策,即使含"电力/运输/项目/能源"等词也不展开**:整份政策若属
    价格/定价/交易/结算/成本监审/财务/税费/补贴/统计分类/数据互联/金融(基金/REITs/信贷)/
    资质认定/评审遴选/行政处罚程序等,**不引入实体环境或人群行为的实际改变**,则
    pathways 返回空或极少(≤1)。例:"输配电定价办法""物流数据开放""REITs行业清单"
    "互联网平台价格规则"——这些**不要**穷举到呼吸/睡眠/慢病。**宁缺勿造,假阳性比漏掉更糟。**
1. **多视角全覆盖**:从这些视角各自审视,别漏:{"、".join(LENSES)}。
2. **沿决定因素逐层展开**(可 2–3 级串联):\n{DETERMINANTS}
3. **重视间接路径**:如"货运增加→道路阻隔/噪声→体力活动下降→慢性病";"用地改变→就业/收入→医疗可及性"。
4. **关注脆弱群体**:老人、儿童、户外劳动者、低收入、慢病人群。
5. **每条路径的 chain 必须完整写出各级节点**,不要用"同上""同A1路径"等交叉引用代替具体链路。
   chain 的**第一个节点用提炼后的简短动作**(如"完善污水管网""推进电炉短流程"),
   **不要照抄政策原文长句**——原文逐字摘录请放进 evidence 字段,chain 只放提炼后的机制节点。

只做**健康因果链**,严守以下三条,否则该工具会失真:
6. **不要重复或近似重复**:同一条"行动→决定因素→健康结果"机制只保留**一条**最完整的;
   不要把同一机制换个措辞重复列,也不要为凑数把一条拆成多条近义路径。
7. **只展开健康因果链,不要做政策设计点评**:像"有效期短/不可持续""罚款收入用途""监管是否到位"
   "信息是否公开"这类**不是**经由健康决定因素到健康结果的因果链,**一律不要列**。
8. **strength 从严,宁低不高**:机制确凿、有广泛文献或常识支撑 → "强";一般合理 → "中";
   凡依赖较多假设、量级不明、或属间接猜测 → **"推测"**。绝不可把推测标成"中/强"。
9. 路径**总数从严控制(通常 ≤12 条)**,只列把握高、与健康有实质关联的;**不为求全面而堆砌
   推测性、脆弱群体细分、二/三级猜测路径**。少而准 >> 多而杂;关联弱的政策本就该极少甚至为 0。

每条路径落到下面 10 个结果问题之一(outcome_q 取 1–10):
{_q_block()}

只输出 JSON:{{"pathways":[{{
 "action_id":"A1","chain":["行动→决定因素→…→健康结果 的逐级节点(2–4 个)"],
 "outcome_q":2,"direction":"风险|效益","population":"受影响人群(尤其脆弱群体)",
 "lens":"{LENSES[0]}等四视角之一","strength":"强|中|推测",
 "status":"文档支持|路径库/文献|假设待证","evidence":"文档原文或机制依据(假设则简述)","confidence":0.0
}}]}}
chain 用简短中文节点;strength 反映路径成立的把握;status 标依据来源。"""


def expand_pathways(actions, doc_text, key, fallback=False):
    alist = "\n".join(f"{a['id']}: {a['action']}" for a in actions)
    extra = ("\n注:本政策仅给出主题、未逐条枚举行动。请**直接依据文档内容**展开主要健康影响路径,"
             "把握高的优先。**但仍须先过第 0b 关**:若文档实质属纯价格/定价/交易/结算/成本监审/"
             "财务/税费/统计分类/数据互联/金融(基金/REITs/信托)/资质认定/评审/行政程序类,"
             "**不引入实体环境或人群行为改变**,则照常返回空或极少(≤1),不要因本提示而硬凑;"
             "只有确为建设/产业/能源/生态/环境/交通/住房等实体政策才系统展开若干条。" if fallback else "")
    user = (f"【行动清单】\n{alist}\n\n【评估对象文档(供引用原文)】\n{doc_text[:25000]}\n\n"
            f"请对每个行动展开多视角因果路径。{extra}只输出 JSON。")
    data = _chat_json(_SYS_EXPAND, user, key, max_tokens=6000)
    return data.get("pathways") or []


# ============ ③ 完整性批判 + 小结 ============
_SYS_CRITIC = f"""你是 HIA 完整性审稿人。给定行动与已展开路径,**仅在确有"重要且把握较高"的遗漏时才补充**。

克制铁律(最重要,优先于"求全"):
- **默认不补**。展开阶段通常已足够;**added 多数情况下应为空,确需补充也一般 0–3 条。**
- **绝不为"求全面"补推测性、脆弱群体细分、二/三级猜测、隐含假设路径**——这类正是要避免的假阳性噪声。
- 只补**确实重要、机制把握高(强/中)**的真实遗漏;依赖较多假设的一律**不补**。
- 不重复/近似已有路径;不把政策设计点评(有效期/罚款/监管/信息公开)当路径。
- 若该政策本身与健康关联很弱(行动非健康类),**added 返回空**,summary 如实说明"与健康关联很弱"。

给一段中立、克制、**如实反映关联强弱(不夸大)**的整体研判小结,以及健康影响程度建议。

只输出 JSON:{{"added":[{{"action_id":"A1","chain":[...],"outcome_q":7,"direction":"风险|效益",
 "population":"...","lens":"...","strength":"强|中|推测","status":"...","evidence":"...","confidence":0.0}}],
 "summary":"2–4 句整体研判小结","suggest_level":"很小|轻度|重大","notes":"仍需专家补充核实的关键缺口"}}"""


def critique_augment(actions, pathways, key):
    alist = "\n".join(f"{a['id']}: {a['action']}" for a in actions)
    plist = "\n".join(
        f"- [{p.get('action_id','?')}→Q{p.get('outcome_q','?')}] "
        f"{' → '.join(p.get('chain', []))}({p.get('strength','')})"
        for p in pathways)
    user = (f"【行动】\n{alist}\n\n【已展开路径】\n{plist}\n\n"
            f"请批判补全并给小结,只输出 JSON。")
    data = _chat_json(_SYS_CRITIC, user, key, max_tokens=4000)
    return (data.get("added") or [], str(data.get("summary", "") or "").strip(),
            str(data.get("suggest_level", "") or "").strip(),
            str(data.get("notes", "") or "").strip())


# ============ 规整 ============
def _norm_pathways(raw, action_ids, start=1):
    out, pid = [], start
    valid_act = set(action_ids)
    seen_paths = set()                   # 整条路径去重:LLM 偶把同一条机制逐字重复列多条(灌水)
    for p in raw:
        if not isinstance(p, dict):
            continue
        try:
            q = int(p.get("outcome_q"))
        except Exception:
            continue
        if not (1 <= q <= len(QUESTIONS)):
            continue
        # 归一化 chain:LLM 有时把它返回成字符串(直接遍历会按"字"拆开),
        # 有时单个元素内部自带 → 箭头;统一拆成干净的环节列表。
        raw_chain = p.get("chain") or []
        if isinstance(raw_chain, str):
            raw_chain = [raw_chain]
        chain = []
        seen = set()
        for c in raw_chain:
            for part in re.split(r"\s*(?:→|—>|->|⟶|⇒|⇨)\s*", str(c)):
                part = part.strip()
                if part and part not in seen:    # 去重:LLM 偶尔让链条绕回先前节点
                    seen.add(part)
                    chain.append(part)
        if not chain:
            continue
        sig = tuple(chain)               # 链路完全相同 → 同一条机制重复,只留首条
        if sig in seen_paths:
            continue
        seen_paths.add(sig)
        aid = str(p.get("action_id", "") or "").strip()
        if aid not in valid_act:
            aid = action_ids[0] if action_ids else "A1"
        st = str(p.get("strength", "")).strip()
        st = st if st in STRENGTHS else "推测"
        status = str(p.get("status", "")).strip()
        status = status if status in STATUSES else "假设待证"
        try:
            conf = max(0.0, min(1.0, float(p.get("confidence", 0))))
        except Exception:
            conf = 0.0
        out.append({
            "id": f"P{pid}", "action_id": aid, "chain": chain, "outcome_q": q,
            "direction": ("效益" if str(p.get("direction", "")).strip() == "效益" else "风险"),
            "population": str(p.get("population", "") or "").strip(),
            "lens": str(p.get("lens", "") or "").strip(),
            "strength": st, "status": status,
            "evidence": str(p.get("evidence", "") or "").strip(), "confidence": conf,
        })
        pid += 1
    return out


# ============ 代码确定性聚合到 10 题 ============
def compute_items(pathways):
    """由(已采纳的)路径确定性聚合出 10 题判定。判断阈值在代码里,不交给 LLM。
    规则:有「强/中」且依据非纯假设的路径 → 是;仅「推测/假设待证」路径 → 不知道;无路径 → 否。"""
    items = []
    for q in range(1, len(QUESTIONS) + 1):
        ps = [p for p in pathways if p["outcome_q"] == q]
        firm = [p for p in ps if p["strength"] in ("强", "中") and p["status"] != "假设待证"]
        if firm:
            ans = "是"
            conf = max(p["confidence"] for p in firm) if firm else 0.0
        elif ps:
            ans = "不知道"
            conf = max(p["confidence"] for p in ps)
        else:
            ans = "否"
            conf = 0.0
        gaps = ""
        if ans == "不知道":
            gaps = "需核实:" + ";".join(" → ".join(p["chain"]) for p in ps[:3])
        items.append({"q": q, "answer": ans, "confidence": conf,
                      "n_path": len(ps), "gaps": gaps})
    return items


_SYS_MAP = """你是严格的 HIA 证据匹配审核员。给你一份"证据卡片目录"(每张:编号 C#、题号、
决定因素关键词)和一批"因果路径"(每条:编号 P#、题号、机制链)。为每条路径**只挑机制真正对应
的那一张卡**,挑不到就留空。

铁律(从严,宁缺毋滥):
- **决定因素必须实质对应**:卡片描述的"决定因素"要正是这条路径中间环节涉及的那个因素,
  且作用方向一致。仅仅"同一个题号"绝对不够。
- **机制不沾边一律不选**。反例(都必须留空,不许硬配):
  · "金属/重金属/交通事故/危化品→伤害" 配到 "自杀(限制农药可得性)" 卡 —— 错,不选;
  · "绿地减少/室外空气/VOCs→呼吸病" 配到 "家用固体燃料燃烧""住房潮湿霉菌" 卡 —— 错,不选;
  · "医疗可及性/收入→慢病" 配到 "烟草""酒精" 卡 —— 错,不选。
- **每条路径最多选 1 张最贴切的卡**;只有确有第 2 张同样精准命中同一机制,才可选 2 张。
- 拿不准、只是"题材相近"→ **不选**(cards 留空)。留空是正常且更可取的。
- **只能从目录里选编号,不得编造**。

只输出 JSON:{"assignments":[{"p":0,"cards":[3]}, ...]}(cards 为卡片编号数组,可空)。"""


def map_evidence(pathways, key, timeout=120):
    """LLM 语义匹配:把每条路径映射到证据库里机制**真正对应**的卡片(同题号)。从严匹配,
    宁缺毋滥(机制不沾边宁可留空、标'证据待补');只选库内卡片、锁题号,不编造。就地修改 pathways。"""
    if not pathways:
        return
    cat = "\n".join(hia_evidence.catalog_lines())
    plist = "\n".join(f"P{i} [Q{p['outcome_q']}] " + " → ".join(p["chain"])
                      for i, p in enumerate(pathways))
    user = (f"【证据卡片目录】\n{cat}\n\n【因果路径】\n{plist}\n\n"
            f"请为每条路径匹配机制相关的卡片编号,只输出 JSON。")
    data = _chat_json(_SYS_MAP, user, key, timeout=timeout, max_tokens=4000, temps=(0.2, 0.5))
    for a in (data.get("assignments") or []):
        try:
            pi = int(a.get("p"))
        except Exception:
            continue
        if not (0 <= pi < len(pathways)):
            continue
        p = pathways[pi]
        qof = f"Q{p.get('outcome_q')}"
        existing = list(p.get("cards") or [])
        seen = {tuple(c["sources"]) for c in existing}
        for ci in (a.get("cards") or []):
            try:
                ci = int(ci)
            except Exception:
                continue
            if not (0 <= ci < len(hia_evidence.CARDS)):
                continue
            if hia_evidence.CARDS[ci]["q"] != qof:          # 锁同题号,防跨域错配
                continue
            ref = hia_evidence.card_ref(ci)
            tk = tuple(ref["sources"])
            if tk in seen:
                continue
            seen.add(tk)
            existing.append(ref)
        p["cards"] = existing[:2]                            # 从严:每条最多 2 张最贴切的


def suggest_level_from(items):
    n_yes = sum(1 for x in items if x["answer"] == "是")
    if n_yes >= 4:
        return "重大"
    if n_yes >= 1:
        return "轻度"
    return "很小"


# ============ 编排:一次完整分析(3 次调用)============
def analyze(doc_text, key, project_name="", progress=None):
    """返回 {actions, pathways, items, summary, suggest_level, notes}。
    progress(stage:str) 可选回调,用于页面显示进度。"""
    def _p(s):
        if progress:
            progress(s)
    _p("① 抽取政策行动…")
    actions, health_relevant = extract_actions(doc_text, key, project_name)
    if not health_relevant:
        # 健康关联门槛:模型判定为表彰/任免/纪念/财税/程序性等非健康政策 → 不强造路径
        return {"actions": [], "pathways": [], "items": compute_items([]),
                "summary": "本政策主要为表彰/奖励/任免/纪念/财税分担/行政程序等内容,未引入实体、"
                           "环境或行为改变;与人群健康的关联很弱,无需开展详细健康影响评估。",
                "suggest_level": "很小",
                "notes": "若认为仍有健康关联,请人工补充具体行动后重评。"}
    fallback_used = not actions
    if not actions:
        # 健康相关但未枚举出具体行动 → 用"政策主题"兜底(去掉"关于印发…的通知"壳,避免被当成程序性通知),
        # 由展开阶段据文档原文展开
        topic = project_name or "评估对象"
        m = re.search(r"《(.+?)》", topic)
        if m:
            topic = m.group(1)
        else:
            topic = re.sub(r"^\d+[_\-\s]*", "", topic)
            topic = re.sub(r"^.*?关于", "", topic)
            topic = re.sub(r"(的)?(通知|决定|意见|公告|批复|函)$", "", topic).strip() or topic
        actions = [{"id": "A1", "action": topic, "evidence": ""}]
    aids = [a["id"] for a in actions]

    _p("② 多视角展开因果路径…")
    raw = expand_pathways(actions, doc_text, key, fallback=fallback_used)
    paths = _norm_pathways(raw, aids, start=1)
    # 安全网:已判健康相关、抽出多条具体行动,展开却 0 路径——多为技术细则型减排政策
    # (煤电升级/储能/节能降碳:行动被抽成工程指标,被逐行动门控误杀)。从政策整体重展开一次。
    if not paths and not fallback_used and len(actions) >= 3:
        _p("②b 行动均为技术指标,按政策目标整体重展开…")
        raw = expand_pathways(actions, doc_text, key, fallback=True)
        paths = _norm_pathways(raw, aids, start=1)

    _p("③ 完整性批判与补全…")
    added, summary, level, notes = critique_augment(actions, paths, key)
    paths += _norm_pathways(added, aids, start=len(paths) + 1)
    _p("④ 匹配证据来源…")
    hia_evidence.annotate(paths)            # 仅用确定性关键词匹配(精准、可审计、可复现)
    # 注:已停用 map_evidence() 的 LLM 语义匹配——它会把同题号但机制不相干的卡片
    # 误贴上去(误导,损"有据可查"),且随机难审计。精准 > 召回,匹配不到宁可标"证据待补"。
    # P2:strength=推测 的路径不挂"已核实"证据卡——牵强/猜测路径不该被国标限值卡背书(损可信度)。
    for p in paths:
        if p.get("strength") == "推测" and p.get("cards"):
            p["cards"] = []

    items = compute_items(paths)
    if level not in ("很小", "轻度", "重大"):
        level = suggest_level_from(items)
    return {"actions": actions, "pathways": paths, "items": items,
            "summary": summary, "suggest_level": level, "notes": notes}


# ============ 因果路径图(DOT,供 st.graphviz_chart)============
def _node_id(prefix, text):
    return prefix + "_" + re.sub(r"\W+", "", text)[:24] + str(abs(hash(text)) % 1000)


def _wrap(text, width=7, max_chars=35):
    """标签按字数折行(DOT 用 \\n 换行):节点变窄变方,压缩整图宽度、改善长宽比、放大有效字号。"""
    s = str(text).replace('"', "'").replace("\\", "/")
    if len(s) > max_chars:
        s = s[:max_chars - 1] + "…"
    return "\\n".join(s[i:i + width] for i in range(0, len(s), width))


def build_dot(actions, pathways):
    """把(已采纳的)路径渲染成 DOT:行动(左)→ 决定因素中间节点 → 结果问题(右)。
    同名中间节点合并;边按强度着色,假设待证用虚线。标签折行 + 大字号,避免被缩放后看不清。"""
    act_label = {a["id"]: a["action"] for a in actions}
    lines = ["digraph G {",
             'rankdir=LR; ranksep=0.9; nodesep=0.4; splines=true;',
             'node [fontname="Microsoft YaHei", fontsize=15, shape=box, '
             'style=rounded, margin="0.14,0.08", color="#9AA0A6"];',
             'edge [fontname="Microsoft YaHei", fontsize=12, arrowsize=0.9];']
    used_actions, used_q, det_nodes = set(), set(), {}
    edges = []

    def det_id(label):
        if label not in det_nodes:
            nid = _node_id("D", label)
            det_nodes[label] = nid
        return det_nodes[label]

    for p in pathways:
        col = {"强": "#1B6B3A" if p["direction"] == "效益" else "#C62828",
               "中": "#2E9E5B" if p["direction"] == "效益" else "#E07B39",
               "推测": "#9AA0A6"}[p["strength"]]
        style = "dashed" if p["status"] == "假设待证" else "solid"
        seq = [("A_" + p["action_id"], None)]
        for step in p["chain"]:
            seq.append((det_id(step), step))
        seq.append((f"Q{p['outcome_q']}", None))
        used_actions.add(p["action_id"])
        used_q.add(p["outcome_q"])
        for (n1, _), (n2, _) in zip(seq, seq[1:]):
            edges.append(f'"{n1}" -> "{n2}" [color="{col}", style={style}, penwidth=1.8];')

    # 行动节点(左):绿底加粗
    lines.append('{ rank=source;')
    for aid in used_actions:
        lab = _wrap(f"{aid} " + act_label.get(aid, aid), width=8, max_chars=40)
        lines.append(f'"A_{aid}" [label="{lab}", style="filled,rounded", '
                     f'fillcolor="#EAF7EF", color="#1B6B3A", fontsize=15];')
    lines.append("}")
    # 决定因素中间节点:折行圆角框
    for lab, nid in det_nodes.items():
        lines.append(f'"{nid}" [label="{_wrap(lab)}"];')
    # 结果问题节点(右):绿框加粗、字更大
    lines.append('{ rank=sink;')
    for q in sorted(used_q):
        lines.append(f'"Q{q}" [label="Q{q}\\n{SHORT_Q[q-1]}", style="filled,rounded", '
                     f'fillcolor="#F6FCF8", color="#2E9E5B", fontsize=16, penwidth=1.6];')
    lines.append("}")
    lines += edges
    lines.append("}")
    return "\n".join(lines)


# ============ 初筛表 docx ============
def build_screen_docx(header, items, pathways, level, expert_opinion):
    """生成填好的《健康影响评估初筛表》docx。pathways 为已采纳路径(用于研判依据章节)。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    SONG, HEI = "宋体", "黑体"
    GREEN = RGBColor(0x1B, 0x6B, 0x3A)
    GREY = RGBColor(0x55, 0x55, 0x55)

    def font(run, name=SONG, size=10.5, bold=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def para(text="", size=10.5, bold=False, color=None, align=None, name=SONG, after=4):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        font(p.add_run(text), name=name, size=size, bold=bold, color=color)
        return p

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = SONG
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), SONG)

    para("健康影响评估初筛表", size=18, bold=True, name=HEI, color=GREEN,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    info = [("评估对象名称", header.get("name", "")), ("发布/实施类别", header.get("category", "")),
            ("起草/提交部门", header.get("dept", "")), ("提交人", header.get("submitter", "")),
            ("电话", header.get("phone", "")), ("初筛日期", header.get("screen_date", "")),
            ("初筛方法", header.get("method", "")), ("涉及的相关部门", header.get("related_dept", ""))]
    t0 = doc.add_table(rows=0, cols=2)
    t0.style = "Table Grid"
    for k, v in info:
        cells = t0.add_row().cells
        font(cells[0].paragraphs[0].add_run(k), name=HEI, size=10.5, bold=True)
        font(cells[1].paragraphs[0].add_run(str(v or "")), size=10.5)
    doc.add_paragraph()

    para("健康影响评估应重点关注的问题", size=12, bold=True, name=HEI, after=6)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    for j, h in enumerate(["#", "重点关注的问题", "是", "不知道", "否"]):
        c = t.rows[0].cells[j]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(c.paragraphs[0].add_run(h), name=HEI, size=10.5, bold=True)
    for it in items:
        cells = t.add_row().cells
        font(cells[0].paragraphs[0].add_run(str(it["q"])), size=10.5)
        font(cells[1].paragraphs[0].add_run(QUESTIONS[it["q"] - 1]), size=10.5)
        for k, label in enumerate(ANSWERS):
            cells[2 + k].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            mark = "☑" if it.get("answer") == label else "□"
            font(cells[2 + k].paragraphs[0].add_run(mark), size=11,
                 bold=(it.get("answer") == label))
    doc.add_paragraph()

    # 研判依据:按题列出采纳的因果路径
    para("研判依据 · 健康影响路径(智能分析展开 · 专家核定)", size=12, bold=True, name=HEI, after=6)
    by_q = {}
    for p in pathways:
        by_q.setdefault(p["outcome_q"], []).append(p)
    for it in items:
        q = it["q"]
        ps = by_q.get(q, [])
        para(f"{q}. {SHORT_Q[q-1]} —— 判定:{it.get('answer','')}"
             + (f"(支撑路径 {len(ps)} 条)" if ps else "(无路径)"),
             size=10, bold=True, after=2)
        for p in ps:
            line = (f"   · [{p['strength']}/{p['status']}] " + " → ".join(p["chain"])
                    + (f"|人群:{p['population']}" if p["population"] else ""))
            para(line, size=9.5, color=GREY, after=1)
            if p.get("status") == "文档支持" and p.get("evidence"):
                para(f"     文档依据:{p['evidence']}", size=9, color=GREY, after=1)
            cards = p.get("cards") or []
            if cards:
                for c in cards:
                    tier = c.get("tier", "WHO 资料")
                    vs = "待补强" if c.get("status") == "todo" else "已核实"
                    para(f"     来源[{tier}·{vs}]:{'；'.join(c['sources'])}",
                         size=9, color=GREY, after=1)
                    if c.get("note"):
                        para(f"       要点(概括,以原文为准):{c['note']}",
                             size=9, color=GREY, after=2)
            else:
                para("     ⚠ 证据待补:此条为机制推断,暂无权威来源支撑,需专家补证。",
                     size=9, color=GREY, after=2)
        if it.get("note"):
            para(f"   专家备注:{it['note']}", size=9.5, after=3)

    doc.add_paragraph()
    para("评估专家组意见", size=12, bold=True, name=HEI, after=4)
    para(expert_opinion or "(待专家填写)", size=10.5, after=8)
    levels = ["很小", "轻度", "重大"]
    para("结论:健康影响程度　" + "　".join(("☑" if level == lv else "□") + lv for lv in levels),
         size=11, bold=True, after=10)
    para("专家组长审定签字:____________　日期:____年__月__日", size=10.5, after=4)
    para("参与专家签字:____________　日期:____年__月__日", size=10.5, after=10)
    para("说明:本表由「健康影响评估智能初筛系统」辅助生成——系统借助智能分析技术,基于上传文档"
         "展开健康影响路径与研判草案;路径采纳、判定与签字以专家核定为准,不替代专家判定。",
         size=9, color=GREY)
    para("制度依据:" + hia_evidence.INSTITUTIONAL_BASIS, size=8.5, color=GREY)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
